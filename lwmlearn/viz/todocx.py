# -*- coding: utf-8 -*-
"""
make office WORD report 

:class:`.WordContainer` class to make ``.docx`` report file

Created on Fri Apr 24 14:09:17 2020

@author: Administrator
"""

#%%

import os

import pandas as pd
import numpy as np
import re

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE

from lwmlearn.utilis.read_write import Objs_management

from lwmlearn.utilis.read_write import search_file
from lwmlearn.utilis.utilis import get_flat_list


class WordContainer():
    """
    Atrributes
    -----------
    report_component_list : list of list
        each list contains 5 element:
           item0:
               key name string of object
           item1:
               bool, whether to output this object, if False, output only text
           item2:
               text string, above object
           item3:
               text string, below object
           item4:
               list of subitems. For example, item4=[] means this object has no 
               subitems, item4=[tuple1, tuple2, ...], same tuple object will 
               be output in a tree hiearchy.       
    """
    def __init__(self,
                 docx=None,
                 search_path='.',
                 title_txt=None,
                 title_align='CENTER'):
        '''
        

        Parameters
        ----------
        docx : TYPE, optional
            docx can be either a path to a .docx file (a string) or a 
            file-like object. If docx is missing or None, the built-in 
            default document "template" is loaded.DESCRIPTION. 
            The default is None.
        
        search_path:
            dir_path to search for graph or excel file object

        Returns
        -------
        None.

        '''
        self.get_docx(docx)
        self._report_component = []
        self.set_search_path(search_path)
        if title_txt is not None:
            self._write_title(title_txt, alignment=title_align)

    def _searchfile(self,
                    filename,
                    suffix=None,
                    subfolder=True,
                    include_suffix=False):

        return search_file(filename, self.search_path, suffix, subfolder,
                           include_suffix)

    def _write_title(self, txt, style='Title', alignment='CENTER'):
        """
        

        Parameters
        ----------
        txt : TYPE
            DESCRIPTION.
        style : TYPE, optional
            DESCRIPTION. The default is 'Title'.

        Returns
        -------
        None.

        """
        document = self.document
        p = document.add_paragraph()
        p.alignment = getattr(WD_ALIGN_PARAGRAPH, alignment)
        if style:
            p.style = style
        p.add_run(txt)

    def _split_fmt_txt(self, txt, pattern='({[^{}]*})', **kwargs):
        """
        splitted 'txt' by pattern, then format each {key} with **kwargs

        Parameters
        ----------
        txt : TYPE
            DESCRIPTION.
        pattern : TYPE, optional
            DESCRIPTION. The default is '({\w*})'.
        **kwargs : TYPE
            DESCRIPTION.

        Returns
        -------
        list
            formatted list of text string.

        """

        txt_lst = re.split(pattern, txt)

        return [i.format_map(kwargs) for i in txt_lst]

    def set_search_path(self, search_path):
        """
        

        Parameters
        ----------
        search_path : path
            path to search for report element.

        Returns
        -------
        None.

        """
        self.search_path = search_path

    @property
    def report_component_list(self):
        '''list of 5 element report item
        
        report_component_list : list of tuples
            each tuple contains 5 element:
                item0 - key name string of object
                item1 - bool, whether to output this object, if False, output only text
                item2 - text string, above object
                item3 - text string, below object
                item4 - list, subitems. an empty [] means this object has no 
                subitems, [tuple, tuple2, ...], same tuple object will be output
                in a tree hiearchy.
        '''
        return self._report_component

    @report_component_list.setter
    def report_component_list(self, report_component_list):
        """
        """
        if not isinstance(report_component_list, list):
            raise ValueError('component_lst must be list type')
        if np.ndim(report_component_list) != 2:
            raise ValueError('Dimension of component_lst must be 2')
        if np.shape(report_component_list)[-1] != 5:
            raise ValueError(
                'length of each iterable in component_lst must be 5')

        self._report_component = report_component_list

    def get_docx(self, docx):
        """initialize self.document        
        
        Parameters
        ----------
        docx : TYPE
            name of docx file. Default is None.

        Returns
        -------
        None.

        """

        self.document = Document(docx)
        self.set_styles()
        return self

    def set_styles(self, style_name='Normal', fontsize=12):
        """set style
        

        Parameters
        ----------
        style_name : TYPE, optional
            check MS word style reference. The default is 'Normal'.
        fontsize : TYPE, optional
            fontsize. The default is 12.

        Returns
        -------
        TYPE
            DESCRIPTION.

        """

        document = self.document
        style = document.styles['Normal']
        document.styles['Normal'].font.name = u'仿宋'
        document.styles['Normal']._element.rPr.rFonts.set(
            qn('w:eastAsia'), u'仿宋')
        font = style.font
        font.size = Pt(fontsize)

        return self

    def write_header(self, txt):
        """
        write page header for self.document

        Parameters
        ----------
        txt : TYPE
            DESCRIPTION.

        Returns
        -------
        None.

        """
        section = self.document.sections[0]
        header = section.header
        paragraph = header.paragraphs[0]
        r = paragraph.add_run(txt)
        r.font.size = Pt(10)
        return

    def write_heading(self, level, txt, fontsize=12, style=None):
        """write paragraph heading text
        

        Parameters
        ----------
        level : int
            level of heading. options are [1, 2, 3 ...]
        txt : TYPE
            text of headings.
        fontsize : TYPE, optional
            DESCRIPTION. The default is 12.

        Returns
        -------
        TYPE
            DESCRIPTION.

        """

        document = self.document
        header = document.add_heading('', level=level)
        run = header.add_run(txt)
        if style:
            header.style = style
        else:
            run.font.name = '宋体'
            run.font.color.rgb = RGBColor(0, 83, 133)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(fontsize)
        run.bold = True
        run.italic = False
        return self

    def write_paragraph(self,
                        txt,
                        fontsize=12,
                        style=None,
                        key_style='Intense Emphasis',
                        **kwargs):
        """
        write 'txt' as paragraph. 'txt' could be formatted by **kwargs. For
        instance txt="XXX {a} XXX", kwargs= {'a' : 'variable'}, txt will be 
        formatted as txt.format_map(kwargs), the result is "XXX variable XXX"

        Parameters
        ----------
        txt : TYPE
            list of text to write to docx in sequence.
        fontsize : TYPE, optional
            DESCRIPTION. The default is 12.
        style : TYPE, optional
            style for paragraph. The default is None.
        key_style : TYPE, optional
            style for key word formatted parts. The default is 'Intense Emphasis'.
            if None, no distinguish style for formatted words
            
        kwargs:
            key words to format txt

        Returns
        -------
        None.

        """

        document = self.document
        p = document.add_paragraph()
        if style:
            p.style = style

        # remove "\n"
        txt = re.sub('\n', '', txt)
        txt_list = self._split_fmt_txt(txt, **kwargs)
        for i, item in enumerate(txt_list):
            if i % 2 > 0 and key_style:
                run = p.add_run(item, style=key_style)
                run.font.size = Pt(fontsize)
                run.font.underline = True
                run.font.italic = False
            else:
                run = p.add_run(item)
                run.font.size = Pt(fontsize)
        return

    def write_graph_obj(self,
                        pic_path,
                        width=6,
                        with_header=True,
                        headerlevel=2,
                        txt_above=False,
                        txt_below=False,
                        fontsize=10,
                        output=True,
                        **kwargs):
        """write graph into docx by input `pic_path` path directory
        

        Parameters
        ----------
        pic_path : str
            graph file path.
        width : TYPE, optional
            DESCRIPTION. The default is 6.
        with_header : TYPE, optional
            whether to write graph name as header. The default is True.
        txt_above : TYPE, optional
            whether to write text before graph. The default is False.
            if given, write a text string before graph .
        txt_below : TYPE, optional
            whether to write text after graph. The default is False.
            if given, write a text string below graph .
        fontsize:
            fontsize for written text string. The default is 
        out_put bool : 
            whether to output graph object. The default is True
        kwargs:
            key words to format txt_above & txt_below
            
        Returns
        -------
        TYPE
            DESCRIPTION.

        """
        # write header
        dirs, file_ext = os.path.split(pic_path)
        filename, ext = os.path.splitext(file_ext)
        if with_header:
            self.write_heading(headerlevel, filename)
        # write txt above graph
        if isinstance(txt_above, str):
            self.write_paragraph(txt_above, fontsize, **kwargs)

        if output:
            document = self.document
            width = Inches(width)
            document.add_picture(pic_path, width=width)
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # write txt below graph
        if isinstance(txt_below, str):
            self.write_paragraph(txt_below, fontsize, **kwargs)

        return self

    def write_table(self, df, style='Colorful List Accent 3'):
        """
        

        Parameters
        ----------
        df : data frame
            table data with headers to write.
        style : TYPE, optional
            DESCRIPTION. The default is 'Colorful List Accent 3'.
            see table_style() returned list of options.

        Returns
        -------
        TYPE
            DESCRIPTION.

        """

        document = self.document
        rows, cols = np.shape(df)
        rows += 1
        table = document.add_table(rows=rows, cols=cols, style=style)
        table.autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr_names = df.columns
        for i, cell in enumerate(table.rows[0].cells):
            cell.text = str(hdr_names[i])
        for ndindex, val in np.ndenumerate(df):
            i, j = ndindex
            table.cell(i + 1, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            table.cell(i + 1, j).text = str(val).strip()
        self.write_paragraph('')
        return self

    def write_table_obj(self,
                        df_path,
                        with_header=True,
                        headerlevel=2,
                        txt_above=False,
                        txt_below=False,
                        fontsize=10,
                        output=True,
                        **kwargs):
        """write table into ``docx`` by input ``df_path`` table path directory
        

        Parameters
        ----------
        df_path : TYPE
            path of data file, suffix include ['.xlsx', '.csv'].
        with_header : TYPE, optional
            whether to write table file name as header. The default is True.
        txt_above : TYPE, optional
            whether to write text before table. The default is False.
            if given, write a text string before table .
        txt_below : TYPE, optional
            whether to write text after table. The default is False.
            if given, write a text string below table .
        fontsize:
            fontsize for written text string. The default is 10
        out_put bool : 
            whether to output table object. The default is True
        kwargs:
            key words to format txt_above & txt_below            
        Returns
        -------
        None.

        """

        # header
        dirs, file_ext = os.path.split(df_path)
        filename, suffix = os.path.splitext(file_ext)
        if with_header:
            self.write_heading(headerlevel, filename)

        if isinstance(txt_above, str):
            self.write_paragraph(txt_above, fontsize, **kwargs)
        # read and write table
        if suffix == '.xlsx':
            data = pd.read_excel(df_path, sheet_name=None)
            for df in data.values():
                self.write_table(df)
        if suffix == '.csv':
            data = pd.read_csv(df_path)
            self.write_table(data)
        # write txt below
        if isinstance(txt_below, str):
            self.write_paragraph(txt_below, fontsize, **kwargs)

        return

    def save_docx(self, tofile):
        '''
        save docx to ``tofile``
        '''

        self.document.save(tofile)

    def run_container(self,
                      report_component_list,
                      headtext_level=1,
                      search_path=None,
                      **kwargs):
        """
        run report_component_list to write component to docx file

        Parameters
        ----------

        report_component_list : list of tuples
            each tuple contains 5 element:
                item0 - key name string of object
                item1 - bool, whether to output this object, if False, output only text
                item2 - text string, above object
                item3 - text string, below object
                item4 - list, subitems. an empty [] means this object has no 
                subitems, [tuple, tuple2, ...], same tuple object will be output
                in a tree hiearchy.
                
        headtext_level : TYPE, optional
            the starting level for headings. The default is 1.
            
        **kwargs:
            key words to formate text strings contained in report_component list

        Returns
        -------
        TYPE
            DESCRIPTION.

        """
        if search_path is not None:
            self.set_search_path(search_path)

        self.report_component_list = report_component_list

        for i in report_component_list:
            # key name of object, str
            key = str(i[0])
            # whether to output this object, bool
            is_output = _bool(i[1])
            # text above object
            text_above = i[2]
            # text below object
            text_below = i[3]
            # subitems
            subitems = i[4]
            # file_path searched by key
            filepath = self._searchfile(key)

            # for subitems add_text after object
            if filepath:
                dir_file, suffix = os.path.splitext(filepath)
                if suffix in {'.csv', '.xlsx'}:
                    self.write_table_obj(filepath,
                                         with_header=True,
                                         headerlevel=headtext_level,
                                         output=is_output,
                                         txt_above=text_above,
                                         txt_below=text_below,
                                         **kwargs)
                else:
                    self.write_graph_obj(filepath,
                                         with_header=True,
                                         headerlevel=headtext_level,
                                         output=is_output,
                                         txt_above=text_above,
                                         txt_below=text_below,
                                         **kwargs)
            else:
                self.write_heading(headtext_level, str(key))
                if isinstance(text_above, str):
                    self.write_paragraph(text_above)
                    self.write_paragraph(text_below)

            # add subitems
            if len(subitems) > 0:
                sub_headtext_level = headtext_level + 1
                self.run_container(subitems, sub_headtext_level, **kwargs)
        return self


def _get_files(dirpath, suffix=None, subfolder=False):
    ''' return file dict {filename : file}

    dirpath - str
        - dir_x to traverse
    suffix -->extension name, or list of extension names, egg ['.xlsx', 'csv']
        - to include file extensions, default None, to include all extensions
    subfolder --> bool
        - true to traverse subfolders, False only the given dirpath
    '''
    if subfolder:
        get_dirs = traverse_all_dirs
    else:
        get_dirs = traverse_dir

    rst = {
        k: v
        for k, v in get_dirs(dirpath).items()
        if os.path.splitext(v)[1] in get_flat_list(suffix) or not suffix
    }
    return rst


def traverse_all_dirs(rootDir):
    '''traverse files under rootDir including subfolders
    return
    ----
    dict - {filename : file}
    '''
    file_dict = dict([file, os.path.join(dirpath, file)]
                     for dirpath, dirnames, filenames in os.walk(rootDir)
                     for file in filenames)
    return file_dict


def traverse_dir(rootDir):
    '''traverse files under rootDir not including subfolder
    return
    ----
    dict - {filename : file}
    '''
    file_dict = {}
    for filename in os.listdir(rootDir):
        pathname = os.path.join(rootDir, filename)
        if os.path.isfile(pathname):
            file_dict[filename] = pathname
    return file_dict


def table_style():
    '''
    show table style list

    Returns
    -------
    list
        DESCRIPTION.

    '''
    document = Document()
    styles = document.styles

    return [s.name for s in styles if s.type == WD_STYLE_TYPE.TABLE]


def _bool(x_str):
    '''
    '''
    return not str(x_str).upper() in {'FALSE'}


def report_config(item_list,
                  report_design,
                  filename='report_config.xlsx',
                  search_path=['graph_output', 'table_output'],
                  **kwargs):
    '''output report structure to excel file 
    
    'filename', containing 3
    sheets ['item_list', 'report_design', 'variables'] so that it could be 
    read by interpret_config()
    
    parameters
    -----
    item_list:
        basic report item list
    report_design:
        nested desgin of report items
    filename:
        file to output config excel file
    **kwargs:
        variables to use in formated text
    Returns
    -------
    None.

    '''
    if len(kwargs) > 0:
        df_kwargs = pd.Series(kwargs).reset_index()
        df_kwargs.columns = ['keyname', 'value']
    else:
        df_kwargs = pd.DataFrame()

    df_search = pd.Series(search_path).to_frame('search_path')

    df_list = [
        _report2df(item_list),
        _report2df(report_design), df_kwargs, df_search
    ]
    df_sheet_name = ['item_list', 'report_design', 'variables', 'search_path']
    writer = Objs_management('.')
    writer.write(df_list, filename, sheet_name=df_sheet_name)

    return


def interpret_config(filename):
    '''
    interpret report_list and variables dict from config file "filename"
    
    return
        report_list
        variables dict
    '''
    df_dict = pd.read_excel(filename, sheet_name=None)
    df1 = np.array(df_dict['item_list'])
    for i in df1:
        i[-1] = _eval(i[-1])

    df2 = np.array(df_dict['report_design'])
    for i in df2:
        i[-1] = _eval(i[-1])

    df_kwargs = df_dict['variables'][['keyname', 'value']]
    kwargs = {k: _eval(v) for k, v in dict(df_kwargs.values).items()}

    df_search = df_dict['search_path']
    kwargs.update({"search_path": df_search['search_path'].tolist()})

    report_list, variable_dict = _get_subitem(df1, df2), kwargs

    return report_list, variable_dict


def _eval(x_str):
    '''
    try eval x_str
    '''

    try:
        return eval(x_str)
    except:
        return x_str


def _report2df(rep_item_list):
    '''
    to make report list as dataframe format
    '''
    new_list = []
    for item in rep_item_list:
        new_list.append([
            item[0],
            item[1],
            item[2],
            item[3],
            [i[0] for i in item[-1]],
        ])
    df_item_list = pd.DataFrame(new_list,
                                columns=[
                                    'key_name', 'output_bool', 'txt_above',
                                    'txt_below', 'subitems'
                                ])
    return df_item_list


def _get_subitem(item_list, report_design):
    '''
    get item value for subitems in report_design list
    '''
    item_d = {i[0]: i for i in item_list}

    def _update_subitems(itm):
        '''
        '''
        itm[-1] = [_update_subitems(item_d[i]) for i in itm[-1]]
        return itm

    return [_update_subitems(i) for i in report_design]


def _retrieve_name(var):
    """
        Gets the name of var. Does it from the out most frame inner-wards.
        :param var: variable to get name from.
        :return: string
        """
    import inspect
    for fi in reversed(inspect.stack()):
        names = [
            var_name for var_name, var_val in fi.frame.f_locals.items()
            if var_val is var
        ]
        if len(names) > 0:
            return names[0]


if __name__ == '__main__':
    pass

    path = [
        r'D:\GitHub\Fintech_viz\fintech_cat\graph_output',
        r'D:\GitHub\Fintech_viz\fintech_cat\table_output'
    ]
    d = WordContainer(search_path=path, title_txt='title', title_align='LEFT')
    ob1 = ("流程监控图1", True, 'above this is a {a9}', 'below', [])
    ob2 = ("业务统计表1", True, 'above', 'below', [])
    report = [
        ["流程监控图1", True, 'above', 'below', [ob1, ob2]],
        ["业务统计表1", True, 'above', 'below', [ob1, ob2]],
    ]
    d.report_component_list = report
    d.run_container(report, a9="(a variable)")
    d.write_header("headers show here")
    d.write_paragraph('Mean ROC(AUC=%0.2f ± %0.2f)' % (1, 0.1))
    d.save_docx('123.docx')
